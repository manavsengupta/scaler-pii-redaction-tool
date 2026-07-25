import re
import os
import io
import tempfile
from datetime import datetime
import pandas as pd
import streamlit as st
from docx import Document
from faker import Faker
from presidio_analyzer import AnalyzerEngine, Pattern, PatternRecognizer, RecognizerRegistry
from presidio_analyzer.nlp_engine import NlpEngineProvider
from names_dataset import NameDataset

st.set_page_config(page_title="PII Redaction Tool", page_icon="🔒", layout="wide")

class Redactor:
    def __init__(self):
        self.fk=Faker('en_US')
        Faker.seed(42)
        self.mp={}
        self.audit_log=[]
        
        # load global name db (takes a couple secs to init in memory)
        print("[~] loading names-dataset corpus...")
        self.nd = NameDataset()
        
        self.reg = RecognizerRegistry()
        self.reg.load_predefined_recognizers()
        self.add_patts()
        
        cfg={"nlp_engine_name": "spacy", "models": [{"lang_code": "en", "model_name": "en_core_web_lg"}]}
        prv=NlpEngineProvider(nlp_configuration=cfg)
        self.anlz=AnalyzerEngine(registry=self.reg, nlp_engine=prv.create_engine())
        
        # spacy instance for grammatical pos checks
        self.nlp = self.anlz.nlp_engine.get_nlp("en")
        
        # domain ignore terms so financial metrics and headers survive
       # massive shield of ignore terms to stop false positives on headers, legal docs, finance & IT words
        self.ign_terms = {
            # currencies & units
            "usd", "eur", "sek", "inr", "rs", "rupees", "dollar", "dollars", "euro", "euros", "krona", 
            "cents", "paise", "lakh", "lakhs", "crore", "crores", "million", "billion", "trillion", 
            "mt", "metric ton", "sqft", "sq mtr", "kwh", "gw", "gwh", "mw", "mva", "kv", "lt", "ht",

            # regulatory & statutory bodies
            "sebi", "rbi", "mca", "roc", "bse", "nse", "cdsl", "nsdl", "rdss", "uidai", "dgft", "dpiit", 
            "cerc", "fbil", "npci", "sbi", "hdfc", "icici", "ibbi", "irdai", "pfrda", "nhai", "mpcb", "msedcl",
            "iatf", "iso", "ul", "bis", "iea", "imf", "weo", "oem", "oems",

            # financial, accounting & valuation terms
            "cagr", "ebitda", "ebit", "pat", "roce", "roe", "rodep", "meis", "epcg", "pli", "kpi", 
            "gaap", "ind as", "ifrs", "us gaap", "indian gaap", "ind as 24", "net worth", "nav", 
            "gross proceeds", "net proceeds", "working capital", "cash flow", "balance sheet", 
            "profit and loss", "p&l", "assets", "liabilities", "depreciation", "amortization", 
            "taxation", "income tax", "gst", "stt", "customs duty", "basic custom duty", "turnover",
            "paid-up", "share capital", "equity shares", "face value", "fresh issue", "bonus issue",
            "split of equity shares", "dividend", "reserves and surplus", "trade receivables", "trade payables",

            # corporate & legal document names
            "prospectus", "red herring prospectus", "drhp", "rhp", "ipo", "fpo", "offer for sale", "ofs", 
            "book building", "price band", "floor price", "cap price", "cut-off price", "bid lot", "offer price",
            "asba", "upi", "can", "allotment advice", "escrow account", "refund account", "public offer account",
            "memorandum of association", "moa", "articles of association", "aoa", "power of attorney", "poa", 
            "certificate of incorporation", "board resolution", "special resolution", "ordinary resolution", 
            "annual return", "audit report", "valuation report", "consent letter", "agreement", "contract", 
            "deed", "indemnity", "undertaking", "affidavit", "notice", "agenda", "minutes", "terms of the offer",
            "offer structure", "general information", "capital structure", "industry overview", "risk factors",
            "forward-looking statements", "definitions and abbreviations", "summary of the offer",

            # corporate roles, departments & entities (when standalone)
            "board of directors", "board", "committee", "audit committee", "ipo committee", 
            "stakeholders relationship committee", "nomination and remuneration committee", "csr committee", 
            "risk management committee", "chairman", "managing director", "whole-time director", 
            "executive director", "independent director", "chief executive officer", "ceo", 
            "chief financial officer", "cfo", "company secretary", "compliance officer", "kmp", 
            "senior management", "promoter", "promoter group", "selling shareholder", "underwriter", 
            "syndicate member", "book running lead manager", "brlm", "registrar to the offer", 
            "statutory auditor", "internal auditor", "secretarial auditor", "legal counsel", "promoters",
            "private limited", "public limited", "public limited company", "private limited company", 
            "limited liability partnership", "llp", "company", "corporation", "limited", "ltd",

            # form labels, contact details & incident headers
            "email", "e-mail", "email address", "phone", "phone number", "telephone", "tel", "mobile", "mob", 
            "fax", "website", "web site", "url", "ip address", "mac address", "device encryption", 
            "encryption key", "operating system", "os", "firmware", "property tag", "serial number", 
            "ticket", "order", "form", "section", "chapter", "clause", "rule", "regulation", "page", 
            "table", "note", "annexure", "schedule", "appendix", "exhibit", "part", "volume", "issue", 
            "version", "date", "timestamp", "description", "short description", "brief synopsis", 
            "in-depth synopsis", "initial assessment", "triage", "usage", "principals", "supervisor", 
            "department", "office location", "job title", "availability", "submission", "revision history",
            "contact person", "registered office", "corporate office", "details", "particulars",

            # IT security & incident terms
            "malware", "phishing", "ransomware", "attachment", "browser", "payload", "shared storage", 
            "usb drive", "network share", "anti-virus", "runtime protection", "periodic scan", "infestation", 
            "quarantine", "unauthorized access", "data breach", "credentials", "login", "password", "pin", 
            "certificate", "private key", "public key", "encryption", "decryption", "firewall", "gateway", 
            "server", "client", "endpoint", "router", "switch", "ethernet", "vpn", "wireless", "wi-fi", 
            "isp", "public network", "server logs", "client logs", "audit trails", "browser history",

            # common nouns, verbs & prepositions that trip up title-case NER
            "for example", "example", "examples", "was detected", "detected", "detection", "of time", 
            "time", "timeline", "timeframe", "on containment", "containment", "miscellaneous", "misc", 
            "photograph", "photo", "image", "curt", "brief", "general", "specific", "total", "aggregate", 
            "average", "weighted average", "approximate", "approx", "maximum", "max", "minimum", "min", 
            "balance", "remainder", "overview", "summary", "introduction", "background", "history", 
            "status", "remarks", "comments", "notes", "references", "others", "nil", "na", "n.a.", 
            "none", "applicable", "not applicable", "yes", "no", "true", "false", "item", "items"
        }

    def norm_k(self, txt):
        cl=re.sub(r"^[\s*#\.\-,():]+|[\s*#\.\-,():]+$", "", txt).strip().lower()
        return re.sub(r"\s+", " ", cl)

    def add_patts(self):
        p1=Pattern(name="in_phone", regex=r"(?<!\d)(?:\+91[\-\s]?)?[6-9]\d{9}(?!\d)", score=0.85)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="IN_PHONE", patterns=[p1]))
        p2=Pattern(name="aadhaar", regex=r"\b\d{4}\s?\d{4}\s?\d{4}\b", score=0.85)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="IN_AADHAAR", patterns=[p2]))
        p3=Pattern(name="pan", regex=r"\b[A-Z]{5}[0-9]{4}[A-Z]{1}\b", score=0.95)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="IN_PAN", patterns=[p3]))
        p4=Pattern(name="ssn", regex=r"\b\d{3}-\d{2}-\d{4}\b", score=0.95)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="US_SSN", patterns=[p4]))
        p5=Pattern(name="ipv4", regex=r"\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b", score=0.90)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="IP_ADDRESS", patterns=[p5]))
        p6=Pattern(name="dob", regex=r"\b(0[1-9]|[12][0-9]|3[01])[- /.](0[1-9]|1[012])[- /.](19|20)\d\d\b", score=0.80)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="DATE_OF_BIRTH", patterns=[p6]))
        
        # non-greedy company regex (stops it from swallowing 15-word sentences ending in Limited!)
        p7=Pattern(name="org_ltd", regex=r"\b(?:[A-Z0-9][a-zA-Z0-9&_\.\-]+)(?:\s+(?:[A-Z0-9][a-zA-Z0-9&_\.\-]+|and|of|for|&|AND|OF|FOR)){0,6}\s+(?:[Ll]imited|LIMITED|[Ll]td\.|LTD\.|LLP|[Pp]rivate\s+[Ll]imited|PRIVATE\s+LIMITED|[Pp]vt\.?\s*[Ll]td\.?|[Cc]orporation|CORPORATION|[Ii]nc\.?|INC\.?)\b", score=0.95)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="ORG", patterns=[p7]))
        
        # pattern to anchor indian 6-digit pin codes
        p8=Pattern(name="in_pin", regex=r"\b[1-8]\d{5}\b", score=0.80)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="IN_PINCODE", patterns=[p8]))

    def is_real_person(self, raw_str):
        # 1. first check our offline name db (using 'or {}' so None values dont crash)
        tokens = re.findall(r"\b[a-zA-Z]+\b", raw_str)
        if tokens:
            for tok in tokens:
                search_res = self.nd.search(tok) or {}
                fn_dict = search_res.get("first_name") or {}
                ln_dict = search_res.get("last_name") or {}
                
                in_fn = (fn_dict.get("country") or {}).get("IN", 0) > 0.001
                in_ln = (ln_dict.get("country") or {}).get("IN", 0) > 0.001
                us_fn = (fn_dict.get("country") or {}).get("US", 0) > 0.005
                us_ln = (ln_dict.get("country") or {}).get("US", 0) > 0.005
                
                if(in_fn or in_ln or us_fn or us_ln):
                    return True # definitely a human name!

        # 2. fallback to spacy pos checking to drop verbs/prepositions
        doc=self.nlp(raw_str)
        for t in doc:
            if(t.is_stop or t.pos_ in ["VERB", "ADP", "DET", "CCONJ", "PRON", "AUX"]):
                return False
        has_prop = any(t.pos_ == "PROPN" or t.shape_.startswith("X") for t in doc)
        return has_prop

    def get_val(self, etype, txt):
        nk=self.norm_k(txt)
        if(nk in self.mp):
            return self.mp[nk]
            
        if(etype in ["ORG", "PERSON"] and len(nk)>4):
            for ex_k, ex_val in self.mp.items():
                if(nk in ex_k or ex_k in nk):
                    self.mp[nk]=ex_val
                    return ex_val

        res=txt
        if(etype in ["PERSON", "PERSON_NAME"]):
            res=self.fk.name()
            if(txt.isupper()):
                res=res.upper()
        elif(etype=="EMAIL_ADDRESS"):
            res=self.fk.email()
        elif(etype in ["PHONE_NUMBER", "IN_PHONE"]):
            res="+91 "+str(self.fk.msisdn()[3:13])
        elif(etype in ["ORG", "COMPANY"]):
            res=self.fk.company()+" Ltd."
            if(txt.isupper()):
                res=res.upper()
        elif(etype in ["LOCATION", "ADDRESS", "IN_ADDRESS"]):
            res=f"{self.fk.building_number()}, {self.fk.street_name()}, {self.fk.city()} - {self.fk.postcode()}"
        elif(etype=="IN_PINCODE"):
            # native check for valid indian postal range (starts with 1-8, exactly 6 digits)
            if(re.match(r"^[1-8]\d{5}$", txt)):
                res = str(self.fk.random_int(110001, 800001))
            else:
                res=txt
        elif(etype=="US_SSN"):
            res=self.fk.ssn()
        elif(etype=="IN_AADHAAR"):
            res="[Aadhaar Redacted]"
        elif(etype=="IN_PAN"):
            res=self.fk.lexify('?????').upper()+self.fk.numerify('####')+self.fk.lexify('?').upper()
        elif(etype=="CREDIT_CARD"):
            res=self.fk.credit_card_number()
        elif(etype=="DATE_OF_BIRTH"):
            res=self.fk.date_of_birth(tzinfo=None, minimum_age=18, maximum_age=80).strftime("%d/%m/%Y")
        elif(etype=="IP_ADDRESS"):
            res=self.fk.ipv4()
        else:
            res=f"[{etype}]"
            
        self.mp[nk]=res
        return res

    def proc_txt(self, txt):
        if not txt.strip():
            return txt
            
        ent_types=["PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER", "IN_PHONE", "ORG", "LOCATION", "US_SSN", "IN_AADHAAR", "IN_PAN", "CREDIT_CARD", "DATE_OF_BIRTH", "IP_ADDRESS", "IN_PINCODE"]
        res_normal=self.anlz.analyze(text=txt, entities=ent_types, language="en")
        
        # shadow parse title-case only for all caps text
        res_shadow=[]
        if(txt.isupper() and len(txt.split())>1):
            res_shadow=self.anlz.analyze(text=txt.title(), entities=["PERSON", "ORG"], language="en")
        
        all_res=res_normal+res_shadow
        all_res.sort(key=lambda x: (x.start, -x.end))
        
        flt=[]
        last_end=-1
        ign_ctx=["ticket", "order", "iso", "form", "section", "regulation", "cin", "iatf", "clause", "rule", "page", "table", "version"]
        
        for r in all_res:
            if(r.start<last_end):
                continue
            old=txt[r.start:r.end]
            nk=self.norm_k(old)
            
            if(nk in self.ign_terms):
                continue
            if(len(nk)<=2 and r.entity_type not in ["IN_PAN", "IN_PHONE"]):
                continue
                
            # upgraded name verification: checks names-dataset FIRST, then spacy pos
            if(r.entity_type=="PERSON" and not self.is_real_person(old)):
                continue
                
            win=txt[max(0, r.start-20):min(len(txt), r.end+20)].lower()
            if(any(w in win for w in ign_ctx) and r.entity_type in ["DATE_OF_BIRTH", "IN_PHONE", "US_SSN", "IN_PINCODE"]):
                continue
                    
            flt.append(r)
            last_end=r.end
            
        flt.sort(key=lambda x: x.start, reverse=True)
        for r in flt:
            old=txt[r.start:r.end]
            new_v=self.get_val(r.entity_type, old)
            
            self.audit_log.append({
                "Timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "Entity Type": r.entity_type,
                "Original Text": old,
                "Synthetic Replacement": new_v,
                "Confidence Score": round(r.score, 2)
            })
            
            txt=txt[:r.start]+new_v+txt[r.end:]
        return txt

    def run_docx(self, inp, outp):
        doc=Document(inp)
        for p in doc.paragraphs:
            if(p.text):
                p.text=self.proc_txt(p.text)
        for t in doc.tables:
            for row in t.rows:
                for c in row.cells:
                    for p in c.paragraphs:
                        if(p.text):
                            p.text=self.proc_txt(p.text)
        doc.save(outp)


# ui layout
st.title("🔒 PII Redaction Tool")
st.subheader("Manav Sengupta - Scaler AI Labs Assignment")
st.caption("upload docx file below to clean private info and download audit logs")

if 'done' not in st.session_state:
    st.session_state.done=False

up_file = st.file_uploader("Select DOCX File", type=["docx"])

if up_file is not None:
    if 'last_file' not in st.session_state or st.session_state.last_file != up_file.name:
        st.session_state.done=False
        st.session_state.last_file=up_file.name

    st.info("Loaded file: "+str(up_file.name))
    
    if not st.session_state.done:
        if st.button("Start Redaction", type="primary"):
            with st.spinner("cleaning document wait..."):
                with tempfile.TemporaryDirectory() as tmp_dir:
                    inp = os.path.join(tmp_dir, "input.docx")
                    outp = os.path.join(tmp_dir, "redacted_output.docx")
                    
                    with open(inp, "wb") as f:
                        f.write(up_file.getbuffer())
                    
                    red = Redactor()
                    red.run_docx(inp, outp)
                    
                    with open(outp, "rb") as f:
                        st.session_state.docx_res = f.read()
                    
                    df_a = pd.DataFrame(red.audit_log)
                    if not df_a.empty:
                        df_a = df_a.drop_duplicates(subset=["Original Text", "Entity Type"])
                    st.session_state.df_a = df_a
                    
                    c_buf = io.StringIO()
                    df_a.to_csv(c_buf, index=False)
                    st.session_state.csv_res = c_buf.getvalue().encode('utf-8')
                    
                    st.session_state.done=True
                    st.rerun()

    if st.session_state.done:
        st.success("done processing!")
        
        c1, c2 = st.columns(2)
        with c1:
            st.download_button(
                label="📄 Download Redacted DOCX",
                data=st.session_state.docx_res,
                file_name="redacted_output.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        with c2:
            st.download_button(
                label="📊 Download Audit Report (CSV)",
                data=st.session_state.csv_res,
                file_name="audit_report.csv",
                mime="text/csv",
                use_container_width=True
            )
        
        st.write("")
        if st.button("🔄 Upload New File / Reset"):
            st.session_state.done=False
            st.rerun()

        if not st.session_state.df_a.empty:
            st.subheader("Redaction Summary")
            m1, m2 = st.columns(2)
            m1.metric("Entities Sanitized", len(st.session_state.df_a))
            m2.metric("Entity Types Found", st.session_state.df_a["Entity Type"].nunique())
            
            st.dataframe(st.session_state.df_a, use_container_width=True)
