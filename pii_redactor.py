import re
import os
import pandas as pd
from datetime import datetime
from docx import Document
from faker import Faker
from presidio_analyzer import AnalyzerEngine, Pattern, PatternRecognizer, RecognizerRegistry
from presidio_analyzer.nlp_engine import NlpEngineProvider

class Redactor:
    def __init__(self):
        self.fk = Faker('en_US')
        Faker.seed(42)
        self.mp = {}
        self.audit_log = []
        
        self.reg = RecognizerRegistry()
        self.reg.load_predefined_recognizers()
        self.add_patts()
        
        cfg = {"nlp_engine_name": "spacy", "models": [{"lang_code": "en", "model_name": "en_core_web_lg"}]}
        prv = NlpEngineProvider(nlp_configuration=cfg)
        self.anlz = AnalyzerEngine(registry=self.reg, nlp_engine=prv.create_engine())
        
        # grab spacy nlp instance directly from presidio to do pos/vocab checks
        self.nlp = self.anlz.nlp_engine.get_nlp("en")
        
        # minimal domain stop words for prospectus headers and financial terms
        self.ign_terms = {
            "usd", "eur", "sek", "inr", "rs", "rupees", "sebi", "icdr", "bse", "nse", "roc",
            "cagr", "ebitda", "pat", "roce", "iso", "iatf", "cin", "equity shares", "face value",
            "fresh issue", "offer price", "floor price", "cap price", "red herring prospectus"
        }

    def norm_k(self, txt):
        cl = re.sub(r"^[\s*#\.\-,():]+|[\s*#\.\-,():]+$", "", txt).strip().lower()
        return re.sub(r"\s+", " ", cl)

    def add_patts(self):
        p1 = Pattern(name="in_phone", regex=r"(?<!\d)(?:\+91[\-\s]?)?[6-9]\d{9}(?!\d)", score=0.85)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="IN_PHONE", patterns=[p1]))
        p2 = Pattern(name="aadhaar", regex=r"\b\d{4}\s?\d{4}\s?\d{4}\b", score=0.85)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="IN_AADHAAR", patterns=[p2]))
        p3 = Pattern(name="pan", regex=r"\b[A-Z]{5}[0-9]{4}[A-Z]{1}\b", score=0.95)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="IN_PAN", patterns=[p3]))
        p4 = Pattern(name="ssn", regex=r"\b\d{3}-\d{2}-\d{4}\b", score=0.95)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="US_SSN", patterns=[p4]))
        p5 = Pattern(name="ipv4", regex=r"\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b", score=0.90)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="IP_ADDRESS", patterns=[p5]))
        p6 = Pattern(name="dob", regex=r"\b(0[1-9]|[12][0-9]|3[01])[- /.](0[1-9]|1[012])[- /.](19|20)\d\d\b", score=0.80)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="DATE_OF_BIRTH", patterns=[p6]))
        p7 = Pattern(name="org_ltd", regex=r"\b[A-Z][a-zA-Z0-9\s]+(?:Limited|Ltd\.|LLP|Private Limited|Pvt\.? Ltd\.?|Corporation|Inc\.?)\b", score=0.95)
        self.reg.add_recognizer(PatternRecognizer(supported_entity="ORG", patterns=[p7]))

    def is_false_person(self, raw_str):
        # check via spacy tokens if text is actually common nouns/verbs/stop words
        doc = self.nlp(raw_str)
        
        # if any token is a stop word or common preposition/verb, it's not a real person name
        for t in doc:
            if t.is_stop or t.pos_ in ["VERB", "ADP", "DET", "CCONJ", "PRON", "AUX"]:
                return True
        
        # if none of the tokens are proper nouns or capitalised words, skip
        has_prop = any(t.pos_ == "PROPN" or t.shape_.startswith("X") for t in doc)
        if not has_prop:
            return True
            
        return False

    def get_val(self, etype, txt):
        nk = self.norm_k(txt)
        if nk in self.mp:
            return self.mp[nk]
            
        # substring check so variations of company/name stay consistent
        if etype in ["ORG", "PERSON"] and len(nk) > 4:
            for ex_k, ex_val in self.mp.items():
                if nk in ex_k or ex_k in nk:
                    self.mp[nk] = ex_val
                    return ex_val

        res = txt
        if etype in ["PERSON", "PERSON_NAME"]:
            res = self.fk.name()
            if txt.isupper():
                res = res.upper()
        elif etype == "EMAIL_ADDRESS":
            res = self.fk.email()
        elif etype in ["PHONE_NUMBER", "IN_PHONE"]:
            res = "+91 " + str(self.fk.msisdn()[3:13])
        elif etype in ["ORG", "COMPANY"]:
            res = self.fk.company() + " Ltd."
            if txt.isupper():
                res = res.upper()
        elif etype in ["LOCATION", "ADDRESS", "IN_ADDRESS"]:
            res = f"{self.fk.building_number()}, {self.fk.street_name()}, {self.fk.city()} - {self.fk.postcode()}"
        elif etype == "US_SSN":
            res = self.fk.ssn()
        elif etype == "IN_AADHAAR":
            res = f"{self.fk.random_int(1000,9999)} {self.fk.random_int(1000,9999)} {self.fk.random_int(1000,9999)}"
        elif etype == "IN_PAN":
            res = self.fk.lexify('?????').upper() + self.fk.numerify('####') + self.fk.lexify('?').upper()
        elif etype == "CREDIT_CARD":
            res = self.fk.credit_card_number()
        elif etype == "DATE_OF_BIRTH":
            res = self.fk.date_of_birth(tzinfo=None, minimum_age=18, maximum_age=80).strftime("%d/%m/%Y")
        elif etype == "IP_ADDRESS":
            res = self.fk.ipv4()
        else:
            res = f"[{etype}]"
            
        self.mp[nk] = res
        return res

    def proc_txt(self, txt):
        if not txt.strip():
            return txt
            
        ent_types = ["PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER", "IN_PHONE", "ORG", "LOCATION", "US_SSN", "IN_AADHAAR", "IN_PAN", "CREDIT_CARD", "DATE_OF_BIRTH", "IP_ADDRESS"]
        res_normal = self.anlz.analyze(text=txt, entities=ent_types, language="en")
        
        # only run title-case shadow parser if text is strictly ALL-CAPS (prevents "for example" bugs)
        res_shadow = []
        if txt.isupper() and len(txt.split()) > 1:
            res_shadow = self.anlz.analyze(text=txt.title(), entities=["PERSON", "ORG"], language="en")
        
        all_res = res_normal + res_shadow
        all_res.sort(key=lambda x: (x.start, -x.end))
        
        flt = []
        last_end = -1
        ign_ctx = ["ticket", "order", "iso", "form", "section", "regulation", "cin", "iatf", "clause", "rule", "page", "table", "version"]
        
        for r in all_res:
            if r.start < last_end:
                continue
            old = txt[r.start:r.end]
            nk = self.norm_k(old)
            
            # domain term filter
            if nk in self.ign_terms:
                continue
            if len(nk) <= 2 and r.entity_type not in ["IN_PAN", "IN_PHONE"]:
                continue
                
            # dynamic spacy pos check to kill false person names
            if r.entity_type == "PERSON" and self.is_false_person(old):
                continue
                
            # operational context check
            win = txt[max(0, r.start-20):min(len(txt), r.end+20)].lower()
            if any(w in win for w in ign_ctx) and r.entity_type in ["DATE_OF_BIRTH", "IN_PHONE", "US_SSN"]:
                continue
                    
            flt.append(r)
            last_end = r.end
            
        flt.sort(key=lambda x: x.start, reverse=True)
        for r in flt:
            old = txt[r.start:r.end]
            new_v = self.get_val(r.entity_type, old)
            
            self.audit_log.append({
                "Timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "Entity Type": r.entity_type,
                "Original Text": old,
                "Synthetic Replacement": new_v,
                "Confidence Score": round(r.score, 2)
            })
            
            txt = txt[:r.start] + new_v + txt[r.end:]
        return txt

    def run_docx(self, inp, outp):
        doc = Document(inp)
        for p in doc.paragraphs:
            if p.text:
                p.text = self.proc_txt(p.text)
        for t in doc.tables:
            for row in t.rows:
                for c in row.cells:
                    for p in c.paragraphs:
                        if p.text:
                            p.text = self.proc_txt(p.text)
        doc.save(outp)
        self.export_audit()

    def export_audit(self):
        if not self.audit_log:
            return
        df = pd.DataFrame(self.audit_log).drop_duplicates(subset=["Original Text", "Entity Type"])
        filename = f"Redaction_Audit_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        df.to_csv(filename, index=False)

if __name__ == "__main__":
    obj = Redactor()
    obj.run_docx("input_document.docx", "redacted_output.docx")
